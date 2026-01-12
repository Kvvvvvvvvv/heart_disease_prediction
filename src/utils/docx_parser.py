"""
DOCX Parser for Patient Data Extraction
Parses uploaded DOCX files to extract patient information
"""
from docx import Document
import re
import json

class PatientDataParser:
    """Parse patient data from DOCX documents"""
    
    def __init__(self):
        self.field_mappings = {
            # Age variations
            'age': ['age', 'patient age', 'years old', 'y/o', 'yo'],
            # Sex variations
            'sex': ['sex', 'gender', 'male', 'female'],
            # Chest pain
            'cp': ['chest pain', 'cp', 'chest pain type', 'angina'],
            # Blood pressure
            'trestbps': ['resting bp', 'resting blood pressure', 'trestbps', 'bp', 'blood pressure', 'systolic'],
            # Cholesterol
            'chol': ['cholesterol', 'chol', 'total cholesterol', 'serum cholesterol'],
            # Fasting blood sugar
            'fbs': ['fbs', 'fasting blood sugar', 'fasting glucose', 'blood sugar'],
            # Resting ECG
            'restecg': ['resting ecg', 'restecg', 'ecg', 'electrocardiogram', 'resting electrocardiographic'],
            # Max heart rate
            'thalach': ['max heart rate', 'thalach', 'maximum heart rate', 'heart rate', 'hr max'],
            # Exercise angina
            'exang': ['exercise angina', 'exang', 'exercise-induced angina', 'angina on exertion'],
            # ST depression
            'oldpeak': ['st depression', 'oldpeak', 'st segment depression', 'st depression induced'],
            # Slope
            'slope': ['slope', 'st slope', 'peak exercise st segment', 'slope of st'],
            # Major vessels
            'ca': ['major vessels', 'ca', 'number of vessels', 'vessels colored', 'fluoroscopy'],
            # Thalassemia
            'thal': ['thal', 'thalassemia', 'thalassemia type', 'thallium']
        }
    
    def parse_docx(self, docx_file):
        """
        Parse DOCX file and extract patient data
        
        Parameters:
        -----------
        docx_file : file-like object
            Uploaded DOCX file
            
        Returns:
        --------
        patient_data : dict
            Extracted patient data
        errors : list
            List of parsing errors/warnings
        """
        try:
            doc = Document(docx_file)
            text_content = self._extract_text(doc)
            patient_data = {}
            errors = []
            
            # Extract all text
            full_text = text_content.lower()
            
            # Parse each field
            for field, keywords in self.field_mappings.items():
                value = self._extract_field(full_text, field, keywords)
                if value is not None:
                    patient_data[field] = value
                else:
                    errors.append(f"Could not find {field} in document")
            
            return patient_data, errors
            
        except Exception as e:
            return {}, [f"Error parsing document: {str(e)}"]
    
    def _extract_text(self, doc):
        """Extract all text from document"""
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)
    
    def _extract_field(self, text, field, keywords):
        """Extract a specific field value from text"""
        for keyword in keywords:
            # Look for patterns like "age: 50" or "age 50" or "age=50"
            patterns = [
                rf'{keyword}\s*[:=]\s*(\d+\.?\d*)',
                rf'{keyword}\s+(\d+\.?\d*)',
                rf'(\d+\.?\d*)\s+{keyword}',
            ]
            
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    value = float(match.group(1))
                    return self._convert_value(field, value)
        
        # Special handling for categorical fields
        if field == 'sex':
            if 'male' in text or 'm' in text:
                return 1
            elif 'female' in text or 'f' in text:
                return 0
        
        if field == 'cp':
            if 'typical angina' in text or 'typical' in text:
                return 0
            elif 'atypical angina' in text or 'atypical' in text:
                return 1
            elif 'non-anginal' in text or 'non anginal' in text:
                return 2
            elif 'asymptomatic' in text:
                return 3
        
        if field == 'fbs':
            if 'yes' in text or '>120' in text or '> 120' in text:
                return 1
            elif 'no' in text or '≤120' in text or '<=120' in text:
                return 0
        
        if field == 'restecg':
            if 'normal' in text:
                return 0
            elif 'st-t' in text or 'st t' in text or 'abnormality' in text:
                return 1
            elif 'hypertrophy' in text or 'left ventricular' in text:
                return 2
        
        if field == 'exang':
            if 'yes' in text:
                return 1
            elif 'no' in text:
                return 0
        
        if field == 'slope':
            if 'upsloping' in text or 'up sloping' in text:
                return 0
            elif 'flat' in text:
                return 1
            elif 'downsloping' in text or 'down sloping' in text:
                return 2
        
        if field == 'thal':
            if 'normal' in text:
                return 1
            elif 'fixed defect' in text or 'fixed' in text:
                return 2
            elif 'reversible defect' in text or 'reversible' in text:
                return 3
        
        return None
    
    def _convert_value(self, field, value):
        """Convert extracted value to appropriate type"""
        if field in ['age', 'trestbps', 'chol', 'thalach', 'ca']:
            return int(value)
        elif field in ['oldpeak']:
            return float(value)
        elif field in ['sex', 'cp', 'fbs', 'restecg', 'exang', 'slope', 'thal']:
            return int(value)
        return value
    
    def create_template_docx(self):
        """Create a template DOCX file for users to fill"""
        doc = Document()
        doc.add_heading('Patient Information Form', 0)
        doc.add_paragraph('Please fill in the following patient information:')
        doc.add_paragraph('')
        
        fields = [
            ('Age', 'Patient age in years (e.g., Age: 50)'),
            ('Sex', 'Male or Female (e.g., Sex: Male)'),
            ('Chest Pain Type', 'Typical Angina, Atypical Angina, Non-anginal Pain, or Asymptomatic'),
            ('Resting Blood Pressure', 'In mm Hg (e.g., Resting BP: 120)'),
            ('Cholesterol', 'In mg/dL (e.g., Cholesterol: 200)'),
            ('Fasting Blood Sugar', 'Yes if >120 mg/dL, No if ≤120 mg/dL'),
            ('Resting ECG', 'Normal, ST-T wave abnormality, or Left ventricular hypertrophy'),
            ('Maximum Heart Rate', 'In bpm (e.g., Max Heart Rate: 150)'),
            ('Exercise Angina', 'Yes or No'),
            ('ST Depression', 'In mm (e.g., ST Depression: 1.0)'),
            ('Slope', 'Upsloping, Flat, or Downsloping'),
            ('Major Vessels', 'Number 0-3 (e.g., Major Vessels: 0)'),
            ('Thalassemia', 'Normal, Fixed defect, or Reversible defect')
        ]
        
        for field, description in fields:
            doc.add_paragraph(f'{field}: {description}', style='List Bullet')
        
        doc.add_paragraph('')
        doc.add_paragraph('Example format:')
        doc.add_paragraph('Age: 50')
        doc.add_paragraph('Sex: Female')
        doc.add_paragraph('Cholesterol: 200')
        doc.add_paragraph('Resting BP: 120')
        doc.add_paragraph('...')
        
        return doc
